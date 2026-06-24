import os
import re
from typing import Dict, Union

from docx import Document
from docx.shared import Inches

# Définition des constantes globales
CHEMIN_TEXTE = "data/Moby_Dick_Or_The_Whale_by_Herman_Melville.txt"
CHEMIN_IMAGE_TITRE = "data/image_finale.jpg"
CHEMIN_GRAPHIQUE = "distribution_paragraphes.png"
CHEMIN_SORTIE = "data/Rapport_Livre.docx"

TITRE_LIVRE = "Moby Dick; Or, The Whale"
AUTEUR_LIVRE = "Herman Melville"
AUTEUR_RAPPORT = "Quiatol Evan et Tongo Herman"


class DocumentProcessingError(Exception):
    """Exception levée pour les erreurs liées à la génération du rapport.

    Args:
        message (str): Description explicite de l'erreur rencontrée.
        code_erreur (int): Code d'erreur numérique pour le suivi interne.
    """

    def __init__(self, message: str, code_erreur: int) -> None:
        super().__init__(message)
        self.code_erreur = code_erreur


def analyser_donnees_texte(chemin_fichier: str) -> Dict[str, Union[int, float, str]]:
    """Extrait les statistiques textuelles requises pour le rapport.

    Args:
        chemin_fichier: Le chemin système vers le fichier texte source.

    Returns:
        Un dictionnaire contenant les métriques calculées sur le texte.

    Raises:
        DocumentProcessingError: Si le fichier est inaccessible ou mal formaté.
    """
    if not chemin_fichier.endswith(".txt"):
        raise DocumentProcessingError(
            "Le fichier d'entrée doit impérativement être un .txt", 400
        )

    try:
        # Ouverture sécurisée garantissant la libération des ressources mémoire
        with open(chemin_fichier, 'r', encoding='utf-8') as fichier:
            contenu = fichier.read()
    except OSError as erreur_systeme:
        raise DocumentProcessingError(
            f"Échec de l'accès au fichier source : {erreur_systeme}", 404
        )

    # Nettoyage et isolation des paragraphes réels
    paragraphes_bruts = contenu.split('\n\n')
    paragraphes = list(filter(lambda p: p.strip() != "", paragraphes_bruts))
    nombre_paragraphes = len(paragraphes)

    # Validation interne critique pour éviter une division par zéro plus tard
    assert nombre_paragraphes > 0, "Le fichier texte analysé est vide."

    # Comptage des mots via une expression génératrice pour limiter
    mots_par_paragraphe = (len(re.findall(r'\b\w+\b', p)) for p in paragraphes)
    liste_mots = list(mots_par_paragraphe)

    min_mots = min(liste_mots)
    max_mots = max(liste_mots)
    moyenne_mots = sum(liste_mots) / nombre_paragraphes

    # Localisation stricte du chapitre 1 pour cibler l'extraction
    mots_chapitre_1 = 0
    motif_chapitre = re.compile(
        r'CHAPTER 1\.(.*?)CHAPTER 2\.', re.IGNORECASE | re.DOTALL
    )
    chapitre_1_match = motif_chapitre.search(contenu)

    if chapitre_1_match is not None:
        texte_chapitre_1 = chapitre_1_match.group(1)
        mots_chapitre_1 = len(re.findall(r'\b\w+\b', texte_chapitre_1))

    return {
        "nombre_paragraphes": nombre_paragraphes,
        "mots_chapitre_1": mots_chapitre_1,
        "min_mots": min_mots,
        "max_mots": max_mots,
        "moyenne_mots": round(moyenne_mots, 2),
        "source": chemin_fichier
    }


def generer_rapport_word(
        statistiques: Dict[str, Union[int, float, str]],
        chemin_sauvegarde: str
) -> None:
    """Compile les informations et crée un fichier Word structuré.

    Args:
        statistiques: Les données textuelles préalablement extraites.
        chemin_sauvegarde: L'emplacement où enregistrer le document .docx.

    Raises:
        DocumentProcessingError: Si une image requise manque ou erreur d'écriture.
    """
    images_requises = (CHEMIN_IMAGE_TITRE, CHEMIN_GRAPHIQUE)
    for chemin_image in images_requises:
        if not os.path.exists(chemin_image):
            raise DocumentProcessingError(
                f"Ressource graphique manquante : {chemin_image}", 404
            )

    document = Document()

    # Page de titre
    titre = document.add_heading(level=0)
    run_titre = titre.add_run(TITRE_LIVRE)
    run_titre.font.bold = True

    document.add_picture(CHEMIN_IMAGE_TITRE, width=Inches(4.0))

    en_tete_auteur = document.add_heading(level=1)
    run_auteur = en_tete_auteur.add_run(f"Auteur : {AUTEUR_LIVRE}")
    run_auteur.font.italic = True

    en_tete_rapport = document.add_heading(level=2)
    run_rapport = en_tete_rapport.add_run(f"Rédigé par : {AUTEUR_RAPPORT}")
    run_rapport.font.bold = True
    run_rapport.font.italic = True

    # Bascule sur la deuxième page pour la séparation logique
    document.add_page_break()

    # Page de Graphique
    document.add_heading("Distribution des longueurs des paragraphes", level=1)
    document.add_picture(CHEMIN_GRAPHIQUE, width=Inches(5.5))

    paragraphe_description = document.add_paragraph()
    run_intro = paragraphe_description.add_run("Description des données : \n")
    run_intro.font.bold = True

    # Concaténation multi-lignes respectant la limite de 79 caractères
    texte_statistiques = (
        f"L'intrigue étudiée comporte un total de "
        f"{statistiques['nombre_paragraphes']} paragraphes.\n"
        f"Le premier chapitre contient spécifiquement "
        f"{statistiques['mots_chapitre_1']} mots.\n"
        f"En analysant la structure, on note un minimum de "
        f"{statistiques['min_mots']} mots "
        f"et un maximum de {statistiques['max_mots']} mots par paragraphe, "
        f"pour une moyenne globale de {statistiques['moyenne_mots']}.\n"
        f"Source des données : {statistiques['source']}"
    )

    paragraphe_description.add_run(texte_statistiques)

    try:
        document.save(chemin_sauvegarde)
    except OSError as erreur_ecriture:
        raise DocumentProcessingError(
            f"Échec de l'enregistrement Word : {erreur_ecriture}", 500
        )


def main() -> None:
    """Point d'entrée orchestant l'analyse puis la génération du rapport."""
    try:
        donnees = analyser_donnees_texte(CHEMIN_TEXTE)
        generer_rapport_word(donnees, CHEMIN_SORTIE)
        print(f"Succès : Le rapport a été généré sous {CHEMIN_SORTIE}")

    except DocumentProcessingError as erreur_metier:
        print(f"Erreur Métier [{erreur_metier.code_erreur}]: {erreur_metier}")
    except Exception as erreur_inattendue:
        print(f"Une exception non gérée s'est produite : {erreur_inattendue}")


if __name__ == "__main__":
    main()